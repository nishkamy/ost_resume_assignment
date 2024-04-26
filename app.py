# from flask import Flask, request, send_file, render_template, jsonify
# import os
# import re
# import pandas as pd
# from PyPDF2 import PdfReader
# from docx import Document
# import docx2txt
# from io import BytesIO
# import logging
# import shutil

# app = Flask(__name__)

# # Regular expression patterns for mobile numbers and email addresses
# mobile_pattern = re.compile(r'\b(?:\+?91[-\s]?)?\d{5}[-\s]?\d{5}(?:\(\w\))?\b')
# email_pattern = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
# # Flask imports and regular expressions remain the same

# @app.route('/')
# def index():
#     # Serve the HTML file upload page
#     return render_template('upload_folder.html')

# @app.route('/upload_folder', methods=['POST'])
# def upload_folder():
#     try:
#         # Get the uploaded files
#         uploaded_files = request.files.getlist('folder_upload')
        
#         # Check if files are uploaded
#         if not uploaded_files:
#             return jsonify({"error": "No files uploaded."}), 400
        
#         # Create a temporary folder to store the uploaded files
#         temp_folder = 'uploaded_folder'
#         os.makedirs(temp_folder, exist_ok=True)
        
#         # Save the uploaded files to the temporary folder
#         for file in uploaded_files:
#             # Save each file in the temporary folder
#             file_path = os.path.join(temp_folder, file.filename)
#             file.save(file_path)
        
#         # Process the files in the temporary folder and get the Excel file
#         excel_file = process_folder(temp_folder)
        
#         # Clean up the temporary folder after processing
#         shutil.rmtree(temp_folder)
        
#         # Return the Excel file as a response
#         return send_file(excel_file, attachment_filename='results.xlsx', as_attachment=True)
    
#     except Exception as e:
#         # Log the error
#         logging.error(f"An unexpected error occurred: {str(e)}")
#         return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

# def process_folder(folder_path):
#     data = []
#     files_in_folder = os.listdir(folder_path)
    
#     for file_name in files_in_folder:
#         file_path = os.path.join(folder_path, file_name)
        
#         mobile_numbers = []
#         email_addresses = []
#         all_text = ""
        
#         # Ensure the file path is correctly formatted
#         if not os.path.isfile(file_path):
#             logging.error(f"File not found: {file_path}")
#             continue
        
#         # Proceed with processing the file
#         try:
#             if file_name.endswith('.pdf'):
#                 # Handle PDF files
#                 with open(file_path, 'rb') as file:
#                     pdf_reader = PdfReader(file)
#                     for page in pdf_reader.pages:
#                         text = page.extract_text()
#                         all_text += text
#                         mobile_matches = mobile_pattern.findall(text)
#                         email_matches = email_pattern.findall(text)
#                         mobile_numbers.extend(mobile_matches)
#                         email_addresses.extend(email_matches)
            
#             elif file_name.endswith('.docx'):
#                 # Handle DOCX files
#                 doc = Document(file_path)
#                 for paragraph in doc.paragraphs:
#                     text = paragraph.text
#                     all_text += text
#                     mobile_matches = mobile_pattern.findall(text)
#                     email_matches = email_pattern.findall(text)
#                     mobile_numbers.extend(mobile_matches)
#                     email_addresses.extend(email_matches)
            
#             elif file_name.endswith('.doc'):
#                 # Handle DOC files
#                 text = docx2txt.process(file_path)
#                 all_text += text
#                 mobile_matches = mobile_pattern.findall(text)
#                 email_matches = email_pattern.findall(text)
#                 mobile_numbers.extend(mobile_matches)
#                 email_addresses.extend(email_matches)
        
#         except Exception as e:
#             logging.error(f"Error while processing file {file_name}: {str(e)}")
#             continue
        
#         # Append the results to the data list
#         data.append({
#             'File Name': file_name,
#             'All Text': all_text,
#             'Mobile Numbers': ', '.join(mobile_numbers),
#             'Email Addresses': ', '.join(email_addresses)
#         })

#     # Convert data to DataFrame
#     df = pd.DataFrame(data)
    
#     # Create a BytesIO object for the Excel file
#     output = BytesIO()
#     df.to_excel(output, index=False)
#     output.seek(0)
    
#     return output

# if __name__ == '__main__':
#     app.run(debug=True, port=5000)


# from flask import Flask, request, send_file, render_template, jsonify
# import os
# import re
# import pandas as pd
# from PyPDF2 import PdfReader
# from docx import Document
# import docx2txt
# from io import BytesIO
# import logging

# app = Flask(__name__)

# # Regular expression patterns for mobile numbers and email addresses
# mobile_pattern = re.compile(r'\b(?:\+?91[-\s]?)?\d{5}[-\s]?\d{5}(?:\(\w\))?\b')
# email_pattern = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')

# @app.route('/')
# def index():
#     # Serve the HTML file upload page
#     return render_template('upload_folder.html')

# @app.route('/upload_folder', methods=['POST'])
# def upload_folder():
#     try:
#         # Get the uploaded files
#         uploaded_files = request.files.getlist('folder_upload')
#         print(uploaded_files)
        
#         # Check if files are uploaded
#         if not uploaded_files:
#             logging.error("No files uploaded.")
#             return jsonify({"error": "No files uploaded."}), 400
        
#         # Create a temporary folder to store the uploaded files
#         temp_folder = 'uploaded_folder'
#         os.makedirs(temp_folder, exist_ok=True)
#         logging.info(f"Temporary folder created at: {temp_folder}")
        
#         # Save the uploaded files to the temporary folder
#         for file in uploaded_files:
#             print(file)
#             file_path = os.path.join(temp_folder, file.filename)
#             print("file path",file_path)
#             file.save(file_path)
#             print(f"File saved at: {file_path}")
        
#         # Process the files in the temporary folder and get the Excel file
#         excel_file = process_folder(temp_folder)
#         print("Processing completed. Sending Excel file.")
        
#         # Return the Excel file as a response
#         return send_file(excel_file, attachment_filename='results.xlsx', as_attachment=True)
    
#     except Exception as e:
#         logging.error(f"An unexpected error occurred: {str(e)}")
#         return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

# def process_folder(folder_path):
#     data = []
#     logging.info(f"Processing files in folder: {folder_path}")
    
#     files_in_folder = os.listdir(folder_path)
#     logging.info(f"Files found in folder: {files_in_folder}")
    
#     for file_name in files_in_folder:
#         file_path = os.path.join(folder_path, file_name)
#         logging.info(f"Processing file: {file_path}")
        
#         mobile_numbers = []
#         email_addresses = []
#         all_text = ""
        
#         # Ensure the file path is correctly formatted
#         if not os.path.isfile(file_path):
#             logging.error(f"File not found: {file_path}")
#             continue
        
#         # Proceed with processing the file
#         try:
#             if file_name.endswith('.pdf'):
#                 with open(file_path, 'rb') as file:
#                     pdf_reader = PdfReader(file)
#                     for page in pdf_reader.pages:
#                         text = page.extract_text()
#                         all_text += text
#                         mobile_matches = mobile_pattern.findall(text)
#                         email_matches = email_pattern.findall(text)
#                         mobile_numbers.extend(mobile_matches)
#                         email_addresses.extend(email_matches)
            
#             elif file_name.endswith('.docx'):
#                 doc = Document(file_path)
#                 for paragraph in doc.paragraphs:
#                     text = paragraph.text
#                     all_text += text
#                     mobile_matches = mobile_pattern.findall(text)
#                     email_matches = email_pattern.findall(text)
#                     mobile_numbers.extend(mobile_matches)
#                     email_addresses.extend(email_matches)
            
#             elif file_name.endswith('.doc'):
#                 text = docx2txt.process(file_path)
#                 all_text += text
#                 mobile_matches = mobile_pattern.findall(text)
#                 email_matches = email_pattern.findall(text)
#                 mobile_numbers.extend(mobile_matches)
#                 email_addresses.extend(email_matches)
        
#         except Exception as e:
#             logging.error(f"Error while processing file {file_name}: {str(e)}")
#             continue
        
#         # Append the results to the data list
#         data.append({
#             'File Name': file_name,
#             'All Text': all_text,
#             'Mobile Numbers': ', '.join(mobile_numbers),
#             'Email Addresses': ', '.join(email_addresses)
#         })

#     # Convert data to DataFrame
#     df = pd.DataFrame(data)
    
#     # Create a BytesIO object for the Excel file
#     output = BytesIO()
#     df.to_excel(output, index=False)
#     output.seek(0)
    
#     logging.info("Excel file created in memory")
    
#     return output

# if __name__ == '__main__':
#     app.run(debug=True, port=5000)

from flask import Flask, request, send_file, render_template, jsonify
import os
import re
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
import docx2txt
import zipfile
from io import BytesIO
import logging

app = Flask(__name__)

# Regular expression patterns for mobile numbers and email addresses
mobile_pattern = re.compile(r'\b(?:\+?91[-\s]?)?\d{5}[-\s]?\d{5}(?:\(\w\))?\b')
email_pattern = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')

@app.route('/')
def index():
    # Serve the HTML file upload page
    return render_template('upload_folder.html')

@app.route('/upload_folder', methods=['POST'])
def upload_folder():
    try:
        # Get the folder from the uploaded files
        folder = request.files.getlist('folder_upload')
        
        # Check if a folder was uploaded
        if not folder:
            return jsonify({"error": "No folder uploaded."}), 400
        
        # Create a temporary folder to store the uploaded files
        temp_folder = 'uploaded_folder'
        os.makedirs(temp_folder, exist_ok=True)
        
        # Save the uploaded files to the temporary folder
        for file in folder:
            # Extract the file name (without path) from the uploaded file
            filename = os.path.basename(file.filename)
            # Normalize the path
            file_path = os.path.join(temp_folder, filename)
            logging.info(f"Saving file: {file_path}")
            file.save(file_path)
        
        # Call the process_folder function to process the uploaded folder
        output_file = process_folder(temp_folder)
        
        # Send the Excel file as a response
        # return send_file(output_file, attachment_filename='results.xlsx', as_attachment=True)
        return send_file(output_file, download_name='results.xlsx', as_attachment=True)

    
    except Exception as e:
        logging.error(f"An unexpected error occurred: {str(e)}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

def process_folder(folder_path):
    # Create a list to hold data for each file
    data = []
    
    # List all files in the folder
    files_in_folder = os.listdir(folder_path)
    
    # Process PDF, DOC, and DOCX files
    for file_name in files_in_folder:
        # Create the full path to the file
        # Normalize the file path
        file_path = os.path.join(folder_path, file_name)
        logging.info(f"Processing file: {file_path}")
        file_path = os.path.normpath(file_path)
        
        # Initialize lists to store mobile numbers and email addresses found in the file
        mobile_numbers = []
        email_addresses = []
        
        # Initialize a variable to store all text from the file
        all_text = ""
        
        # Try to process the file based on its extension
        try:
            if file_name.endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        all_text += text
                        mobile_matches = mobile_pattern.findall(text)
                        email_matches = email_pattern.findall(text)
                        mobile_numbers.extend(mobile_matches)
                        email_addresses.extend(email_matches)
            
            elif file_name.endswith('.docx'):
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    text = paragraph.text
                    all_text += text
                    mobile_matches = mobile_pattern.findall(text)
                    email_matches = email_pattern.findall(text)
                    mobile_numbers.extend(mobile_matches)
                    email_addresses.extend(email_matches)
            
            elif file_name.endswith('.doc'):
                text = docx2txt.process(file_path)
                all_text += text
                mobile_matches = mobile_pattern.findall(text)
                email_matches = email_pattern.findall(text)
                mobile_numbers.extend(mobile_matches)
                email_addresses.extend(email_matches)
        
        except zipfile.BadZipFile:
            logging.error(f"Error processing file {file_name}: The file is not a valid DOCX file.")
        
        except Exception as e:
            logging.error(f"An error occurred while processing file {file_name}: {str(e)}")
        
        data.append({
            'File Name': file_name,
            'All Text': all_text,
            'Mobile Numbers': ', '.join(mobile_numbers),
            'Email Addresses': ', '.join(email_addresses)
        })

    df = pd.DataFrame(data)
    output = BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    return output

if __name__ == '__main__':
    app.run(debug=True, port=5000)
